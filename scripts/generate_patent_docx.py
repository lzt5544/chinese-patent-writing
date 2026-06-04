#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中国专利申请文件 Word 生成器

符合《专利审查指南》格式要求：
- A4 纸张（210mm × 297mm），纵向单面
- 页边距：上25mm，左25mm，右15mm，下15mm
- 字体：宋体（中文）+ Times New Roman（西文）
- 字高：3.5-4.5mm（小四号 12pt ≈ 4.23mm）
- 行距：约 2.5-3.5mm 行间距离
- 页码：页脚居中
- 说明书不得含有段落编号（如[0001]）
- 权利要求书用阿拉伯数字顺序编号：1., 2., 3.
- 公式：支持 LaTeX 语法（$...$ 标记），自动转换为 Word 原生 Office Math 公式

用法：
    python generate_patent_docx.py <input.json> <output.docx>

input.json 格式：
{
    "patent_type": "发明专利",
    "invention_name": "发明名称",
    "sections": {
        "claims": [
            "权利要求1内容，可包含 $L_{target}$ 这样的公式",
            "权利要求2内容..."
        ],
        "specification": {
            "tech_field": "技术领域内容",
            "background": "背景技术内容",
            "invention_content": {
                "problem": "技术问题",
                "solution": "技术方案内容，公式如 $L_{target}(\\delta) = L_{max} \\cdot (1-\\delta/\\delta_{max})$",
                "effects": "有益效果"
            },
            "figure_desc": "附图说明",
            "embodiment": "具体实施方式"
        },
        "abstract": {
            "text": "摘要正文（≤300字）",
            "figure": "图1"
        }
    }
}

LaTeX 公式说明：
- 用 $...$ 包裹 LaTeX 公式，脚本自动转换为 Word 原生公式
- 支持：下标 _ 、上标 ^ 、分数 \\frac{a}{b} 、根号 \\sqrt{a} 、希腊字母 \\alpha \\beta \\delta 等
- 支持运算符：\\cdot \\times \\pm \\leq \\geq \\neq \\infty \\rightarrow 等
- 示例：$L_{target}(\\delta) = L_{max} \\cdot (1-\\delta/\\delta_{max}) + L_{min} \\cdot (\\delta/\\delta_{max})$
"""

import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误：未安装 python-docx。请运行：pip install python-docx")
    sys.exit(1)


# ==================== 专利格式常量 ====================

PAGE_WIDTH = Cm(21)
PAGE_HEIGHT = Cm(29.7)

TOP_MARGIN = Cm(2.5)
LEFT_MARGIN = Cm(2.5)
RIGHT_MARGIN = Cm(1.5)
BOTTOM_MARGIN = Cm(1.5)

FONT_NAME_CN = '宋体'
FONT_NAME_EN = 'Times New Roman'

FONT_SIZE_TITLE = Pt(16)
FONT_SIZE_NAME = Pt(14)
FONT_SIZE_HEADING1 = Pt(14)
FONT_SIZE_HEADING2 = Pt(12)
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_SMALL = Pt(10.5)

LINE_SPACING_BODY = Pt(18)
LINE_SPACING_TITLE = Pt(24)

FIRST_LINE_INDENT = Cm(0.74)

# Office Math 命名空间
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


# ==================== LaTeX 公式处理 ====================

def latex_to_office_linear(latex):
    """
    将 LaTeX 公式转换为 Office Math linear format。
    Office Math linear format 是 Word 原生公式编辑器支持的类 LaTeX 语法。
    """
    text = latex.strip()

    # 转换 \frac{a}{b} → (a)/(b)
    # 需要递归处理嵌套分数，使用循环而非递归
    prev = None
    while prev != text:
        prev = text
        text = re.sub(
            r'\\frac\{([^{}]*)\}\{([^{}]*)\}',
            lambda m: '(%s)/(%s)' % (m.group(1), m.group(2)),
            text
        )

    # 转换 \sqrt{a} → \sqrt(a)
    text = re.sub(r'\\sqrt\{([^{}]*)\}', r'\\sqrt(\1)', text)

    # 转换 \sqrt[n]{a} → \sqrt(n&a)
    text = re.sub(r'\\sqrt\[(\d+)\]\{([^{}]*)\}', r'\\sqrt(\1&\2)', text)

    # 删除 \left 和 \right（Office 自动调整括号大小）
    text = text.replace(r'\left', '').replace(r'\right', '')

    # 转换省略号
    text = text.replace(r'\dots', '...').replace(r'\ldots', '...')
    text = text.replace(r'\cdots', '⋯')

    # 转换角度
    text = text.replace(r'^∘', '^°')

    return text


def create_omml_formula(office_linear):
    """
    创建 Office Math ML (OMML) 公式元素。
    返回一个 m:oMath 元素，可直接插入到 Word 段落中。
    """
    # 创建 m:oMath 元素
    omath = OxmlElement('m:oMath')

    # 创建 m:r (run) 元素
    mr = OxmlElement('m:r')

    # 创建 m:rPr (run properties) 元素，设置数学字体
    mrpr = OxmlElement('m:rPr')
    msty = OxmlElement('m:sty')
    msty.set(qn('m:val'), 'p')  # p = plain text style
    mrpr.append(msty)
    mr.append(mrpr)

    # 创建 m:t (text) 元素
    mt = OxmlElement('m:t')
    mt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    mt.text = office_linear
    mr.append(mt)

    omath.append(mr)
    return omath


def add_text_with_latex_to_para(para, text):
    """
    向单个段落添加文本，支持 $...$ 标记的 LaTeX 公式。
    不处理 \n（调用者应确保 text 不含 \n）。
    """
    if not text:
        return

    parts = re.split(r'(\$[^$]+\$)', text)
    for part in parts:
        if part.startswith('$') and part.endswith('$') and len(part) > 2:
            latex = part[1:-1]
            office_linear = latex_to_office_linear(latex)
            run = para.add_run()
            run.font.name = FONT_NAME_EN
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)
            omath = create_omml_formula(office_linear)
            run._r.append(omath)
        else:
            if part:
                run = para.add_run(part)
                set_chinese_font(run)


def add_body_paragraphs(doc, text, first_line_indent=True):
    """
    添加正文段落，支持 $...$ LaTeX 公式，支持 \n 拆分为多个段落（硬回车）。
    每个 \n 创建一个独立的段落，绝不使用软回车（Shift+Enter）。
    """
    if not text:
        return

    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        para = doc.add_paragraph()
        add_text_with_latex_to_para(para, line)
        if first_line_indent:
            para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = LINE_SPACING_BODY
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)


def add_claim_with_latex(doc, number, text):
    """添加权利要求项，支持 LaTeX 公式和 \n 分段"""
    lines = text.split('\n')
    for idx, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        para = doc.add_paragraph()

        if idx == 0:
            # 第一行加编号
            run_num = para.add_run(f"{number}.")
            set_chinese_font(run_num)
            if '$' in line:
                run_space = para.add_run(' ')
                set_chinese_font(run_space)
                add_text_with_latex_to_para(para, line)
            else:
                run_text = para.add_run(f" {line}")
                set_chinese_font(run_text)
        else:
            # 后续行缩进对齐（权利要求多行时的缩进）
            if '$' in line:
                add_text_with_latex_to_para(para, line)
            else:
                run_text = para.add_run(line)
                set_chinese_font(run_text)

        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = LINE_SPACING_BODY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.first_line_indent = Pt(0)


# ==================== 工具函数 ====================

def set_chinese_font(run, font_size=None, bold=False):
    """设置中文宋体 + 英文 Times New Roman"""
    if font_size is None:
        font_size = FONT_SIZE_BODY
    run.font.name = FONT_NAME_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_page_number(section):
    """在页脚添加居中页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    set_chinese_font(run, FONT_SIZE_SMALL)


def setup_section(section):
    """设置页面格式（A4 + 专利页边距）"""
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.orientation = WD_ORIENT.PORTRAIT

    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN

    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def setup_normal_style(doc):
    """设置正文默认样式"""
    style = doc.styles['Normal']
    style.font.name = FONT_NAME_EN
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
    style.font.size = FONT_SIZE_BODY
    style.font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_SPACING_BODY
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


# ==================== 内容添加函数 ====================

def add_document_title(doc, text):
    """添加文档大标题"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_chinese_font(run, FONT_SIZE_TITLE, bold=True)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = LINE_SPACING_TITLE
    return para


def add_invention_name(doc, name):
    """添加发明名称"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(name)
    set_chinese_font(run, FONT_SIZE_NAME, bold=True)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = LINE_SPACING_TITLE
    return para


def add_heading1(doc, text):
    """添加一级标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, FONT_SIZE_HEADING1, bold=True)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = LINE_SPACING_BODY
    return para


def add_heading2(doc, text):
    """添加二级标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, FONT_SIZE_HEADING2, bold=True)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = LINE_SPACING_BODY
    return para


def add_body_text_plain(doc, text, first_line_indent=True):
    """添加纯文本段落（无公式），支持 \n 拆分为多个段落（硬回车）"""
    if not text:
        return
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        para = doc.add_paragraph()
        run = para.add_run(line)
        set_chinese_font(run)
        if first_line_indent:
            para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = LINE_SPACING_BODY
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)


# ==================== 主生成函数 ====================

def generate_patent_docx(data, output_path):
    """生成专利 Word 文档"""
    doc = Document()

    # 设置页面和默认样式
    section = doc.sections[0]
    setup_section(section)
    setup_normal_style(doc)
    add_page_number(section)

    # 获取数据
    patent_type = data.get("patent_type", "发明专利")
    invention_name = data.get("invention_name", "未命名发明")
    sections = data.get("sections", {})

    # ============== 权利要求书 ==============
    add_document_title(doc, "权 利 要 求 书")

    claims = sections.get("claims", [])
    for i, claim in enumerate(claims, 1):
        add_claim_with_latex(doc, i, claim)

    # 分页
    doc.add_page_break()

    # ============== 说明书 ==============
    add_document_title(doc, "说 明 书")
    add_invention_name(doc, invention_name)

    spec = sections.get("specification", {})

    # 技术领域
    add_heading1(doc, "技术领域")
    tech_field = spec.get("tech_field", "")
    if tech_field:
        add_body_paragraphs(doc, tech_field)

    # 背景技术
    add_heading1(doc, "背景技术")
    background = spec.get("background", "")
    if background:
        add_body_paragraphs(doc, background)

    # 发明内容
    add_heading1(doc, "发明内容")
    invention_content = spec.get("invention_content", {})

    problem = invention_content.get("problem", "")
    if problem:
        add_heading2(doc, "技术问题")
        add_body_paragraphs(doc, problem)

    solution = invention_content.get("solution", "")
    if solution:
        add_heading2(doc, "技术方案")
        add_body_paragraphs(doc, solution)

    effects = invention_content.get("effects", "")
    if effects:
        add_heading2(doc, "有益效果")
        add_body_paragraphs(doc, effects)

    # 附图说明
    add_heading1(doc, "附图说明")
    figure_desc = spec.get("figure_desc", "")
    if figure_desc:
        add_body_paragraphs(doc, figure_desc)

    # 具体实施方式
    add_heading1(doc, "具体实施方式")
    embodiment = spec.get("embodiment", "")
    if embodiment:
        add_body_paragraphs(doc, embodiment)

    # 分页
    doc.add_page_break()

    # ============== 说明书摘要 ==============
    add_document_title(doc, "说明书摘要")

    abstract_data = sections.get("abstract", {})
    abstract_text = abstract_data.get("text", "")
    if abstract_text:
        add_body_paragraphs(doc, abstract_text)

    abstract_figure = abstract_data.get("figure", "")
    if abstract_figure:
        para = doc.add_paragraph()
        run = para.add_run(f"摘要附图：{abstract_figure}")
        set_chinese_font(run)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = LINE_SPACING_BODY
        para.paragraph_format.space_before = Pt(6)

    # 保存
    doc.save(output_path)
    print(f"专利文档已生成: {output_path}")
    print(f"  - 发明名称: {invention_name}")
    print(f"  - 专利类型: {patent_type}")
    print(f"  - 权利要求数: {len(claims)}")
    print(f"  - 格式: A4纸张, 宋体小四, 专利标准页边距, LaTeX公式支持")


def main():
    """命令行入口"""
    if len(sys.argv) < 3:
        print("用法: python generate_patent_docx.py <input.json> <output.docx>")
        print("")
        print("input.json 应包含专利内容，格式参见脚本顶部注释")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not input_path.exists():
        print(f"错误: 输入文件不存在: {input_path}")
        sys.exit(1)

    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    generate_patent_docx(data, output_path)


if __name__ == '__main__':
    main()
